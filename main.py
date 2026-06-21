import json
import uuid
import os
import gc  # Garbage collector
import torch # GPU memory management
import pythoncom # REQUIRED for MS Word automation
from datetime import datetime
from flask import Flask, request, render_template, send_from_directory, redirect, session, send_file
from werkzeug.utils import secure_filename

# --- PDF & DOCX IMPORTS ---
from pdf2docx import Converter
from docx import Document
from docx2pdf import convert as convert_to_pdf
from deep_translator import GoogleTranslator

# --- IMPORT YOUR HELPER MODULES ---
from audio_sync import AudioSync
from image_translator import ImageTranslator
from audio_from_video import AudioFromVideo
from text_from_audio import TextFromAudio
from text_to_lan import TextToLan
from textlan_to_audio import TextlanToAudio
from translations import TRANSLATIONS

# --- CONFIGURATION ---
BASE_DIR = os.path.abspath(os.path.dirname(__file__))
DB_FILE = os.path.join(BASE_DIR, 'db', 'videos.json') 

# Define Folders
VIDEO_FOLDER = os.path.join(BASE_DIR, 'db', 'video')
AUDIO_FOLDER = os.path.join(BASE_DIR, 'db', 'audio')
AUDIO_ENG_FOLDER = os.path.join(AUDIO_FOLDER, 'eng_audio')
AUDIO_OUTPUT_FOLDER = os.path.join(AUDIO_FOLDER, 'output_au')
IMAGE_FOLDER = os.path.join(BASE_DIR, 'db', 'image')
PDF_INPUT_FOLDER = os.path.join(BASE_DIR, 'db', 'pdf', 'input')
PDF_OUTPUT_FOLDER = os.path.join(BASE_DIR, 'db', 'pdf', 'output')

# Create directories if they don't exist
for folder in [VIDEO_FOLDER, AUDIO_ENG_FOLDER, AUDIO_OUTPUT_FOLDER, IMAGE_FOLDER, PDF_INPUT_FOLDER, PDF_OUTPUT_FOLDER]:
    os.makedirs(folder, exist_ok=True)

# Initialize DB
if not os.path.exists(DB_FILE):
    with open(DB_FILE, 'w') as f: json.dump([], f)

app = Flask(__name__)
app.secret_key = 'hackathon_secret'

# --- PDF LANGUAGE MAP ---
PDF_LANGUAGES = {
    "Assamese": "as", "Bengali": "bn", "Gujarati": "gu", "Hindi": "hi",
    "Kannada": "kn", "Malayalam": "ml", "Marathi": "mr", "Odia": "or",
    "Punjabi": "pa", "Tamil": "ta", "Telugu": "te", "Urdu": "ur"
}

# --- HELPER FUNCTIONS ---
def load_videos():
    try:
        with open(DB_FILE, 'r') as f: return json.load(f)
    except: return []

def save_video_entry(entry):
    videos = load_videos()
    videos.append(entry)
    with open(DB_FILE, 'w') as f: json.dump(videos, f, indent=4)

def get_video_by_id(vid_id):
    videos = load_videos()
    for video in videos:
        if video['id'] == vid_id: return video
    return None

def clear_memory():
    gc.collect()
    if torch.cuda.is_available():
        torch.cuda.empty_cache()

# --- PDF PROCESSING HELPER ---
def process_pdf_file(input_pdf_path, target_lang):
    filename = os.path.basename(input_pdf_path).replace('.pdf', '')
    # Temp DOCX goes to input folder (or a temp folder)
    docx_path = os.path.join(PDF_INPUT_FOLDER, f"{filename}.docx")
    final_pdf_path = os.path.join(PDF_OUTPUT_FOLDER, f"{filename}_{target_lang}.pdf")

    # 1. PDF -> DOCX
    print("   📄 Converting PDF to Word...")
    try:
        cv = Converter(input_pdf_path)
        cv.convert(docx_path, start=0, end=None)
        cv.close()
    except Exception as e:
        print(f"Error in PDF conversion: {e}")
        raise

    # 2. Translate DOCX
    print("   🔠 Translating text...")
    try:
        doc = Document(docx_path)
        
        def translate_text_google(text, lang):
            if not text or len(text.strip()) < 2: return text
            try:
                # Using Google Translator for PDF text as it's robust for general docs
                return GoogleTranslator(source='auto', target=lang).translate(text)
            except: return text

        def process_paragraphs(paragraphs):
            for p in paragraphs:
                for run in p.runs:
                    if run.text.strip():
                        run.text = translate_text_google(run.text, target_lang)
                        # Set a font that supports Indian languages
                        run.font.name = 'Arial Unicode MS'

        process_paragraphs(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    process_paragraphs(cell.paragraphs)

        doc.save(docx_path)
    except Exception as e:
        print(f"Error in translation loop: {e}")
        raise

    # 3. DOCX -> PDF
    print("   🔄 Converting back to PDF...")
    try:
        # Critical for Flask + Word Automation
        pythoncom.CoInitialize() 
        convert_to_pdf(docx_path, final_pdf_path)
    except Exception as e:
        print(f"Error in Word to PDF conversion: {e}")
        raise
    
    return final_pdf_path

# --- CONTEXT PROCESSOR ---
@app.context_processor
def inject_translations():
    def get_text(key):
        lang = session.get('language', 'en')
        return TRANSLATIONS.get(lang, TRANSLATIONS['en']).get(key, key)
    return dict(t=get_text, current_lang=session.get('language', 'en'))

@app.route('/set_language/<lang_code>')
def set_language(lang_code):
    if lang_code in TRANSLATIONS: session['language'] = lang_code
    return redirect(request.referrer or '/')

# --- ROUTES ---

@app.route('/')
def home():
    videos = load_videos()
    return render_template('home.html', videos=videos)

@app.route('/upload', methods=['GET', 'POST'])
def index():
    if request.method == 'POST': return upload()
    return render_template('upload.html')

def upload():
    if 'videoFile' not in request.files: return "No file", 400
    file = request.files['videoFile']
    if file.filename == '': return "No file", 400

    if file:
        # 1. Save Video
        filename = secure_filename(file.filename)
        video_path = os.path.join(VIDEO_FOLDER, filename)
        file.save(video_path)

        # 2. Get Form Data
        title = request.form.get('title', 'Untitled')
        desc = request.form.get('description', '')
        source_lang = request.form.get('sourceLanguage', 'en')
        target_langs = request.form.getlist('targetLanguages') 
        
        if not target_langs: target_langs = ['ta']

        video_id = str(uuid.uuid4())
        
        try:
            print(f"Processing Video: {filename}")

            # 3. Extract Source Audio
            print(" Extracting Audio...")
            ext_audio = AudioFromVideo(video_path, AUDIO_ENG_FOLDER)
            source_audio_path = ext_audio.extracting_audio_eng()

            # 4. Transcribe (Source Audio -> Source Text)
            print(f" Transcribing ({source_lang})...")
            # Uses Whisper (Base Model recommended for 4GB VRAM)
            ext_text = TextFromAudio(source_audio_path, language=source_lang) 
            source_text = ext_text.extracting_text()

            # Clear memory immediately after Whisper
            clear_memory()

            content_map = {}

            # 5. Loop through Target Languages
            for lang in target_langs:
                print(f"🚀 Processing Language: {lang}...")
                
                if lang == source_lang:
                    continue

                lang_data = {}

                # A. Translate
                translator = TextToLan(source_text, source_lang=source_lang, target_lang=lang)
                final_text = translator.convert()
                lang_data['text'] = final_text

                # B. TTS
                tts = TextlanToAudio(final_text, target_lang=lang)
                raw_audio = tts.tamil_audio_conv(AUDIO_OUTPUT_FOLDER, unique_prefix=f"temp_{video_id}")

                # C. Sync
                if raw_audio:
                    final_filename = f"{video_id}_{lang}.wav"
                    final_path = os.path.join(AUDIO_OUTPUT_FOLDER, final_filename)
                    syncer = AudioSync(video_path, raw_audio)
                    syncer.sync_audio(final_path)
                    try: os.remove(raw_audio)
                    except: pass
                    lang_data['audio_file'] = final_filename
                else:
                    lang_data['audio_file'] = None

                content_map[lang] = lang_data
                clear_memory()

            # 6. Save to DB
            new_entry = {
                "id": video_id,
                "title": title,
                "description": desc,
                "original_filename": filename,
                "content_map": content_map,
                "timestamp": datetime.now().strftime("%Y-%m-%d")
            }
            save_video_entry(new_entry)
            
            print("Process Complete!")
            return redirect('/')

        except Exception as e:
            print(f" ERROR: {e}")
            import traceback
            traceback.print_exc()
            return f"Error: {e}", 500

@app.route('/view/<video_id>')
def view_video(video_id):
    video = get_video_by_id(video_id)
    if not video: return "Video not found", 404
    
    return render_template('lmsview.html', 
                           video=video,  
                           content_map=json.dumps(video.get('content_map', {})))

@app.route('/image_translate', methods=['GET', 'POST'])
def image_translate():
    processed_image = None
    original_image = None
    
    if request.method == 'POST':
        if 'imageFile' not in request.files: return "No file", 400
        file = request.files['imageFile']
        
        if file:
            filename = secure_filename(file.filename)
            input_path = os.path.join(IMAGE_FOLDER, filename)
            file.save(input_path)
            
            output_filename = f"translated_{filename}"
            output_path = os.path.join(IMAGE_FOLDER, output_filename)
            
            font_path = os.path.join(BASE_DIR, 'static', 'NotoSansTamil-ExtraBold.ttf')
            translator = ImageTranslator(font_path=font_path)
            
            success = translator.process(input_path, output_path, target_lang='ta')
            
            if success:
                processed_image = output_filename
                original_image = filename
    
    return render_template('image_translate.html', 
                         original=original_image, 
                         processed=processed_image)

# --- NEW: PDF TRANSLATOR ROUTE ---
@app.route('/pdf_translate', methods=['GET', 'POST'])
def pdf_translate():
    if request.method == 'POST':
        file = request.files.get('file')
        target_lang = request.form.get('language')
        target_code = PDF_LANGUAGES.get(target_lang)

        if file and target_code:
            filename = secure_filename(file.filename)
            input_path = os.path.join(PDF_INPUT_FOLDER, filename)
            file.save(input_path)

            try:
                print(f" Starting PDF Translation: {filename} -> {target_lang}")
                output_pdf = process_pdf_file(input_path, target_code)
                print(f" PDF Saved: {output_pdf}")
                return send_file(output_pdf, as_attachment=True)
            except Exception as e:
                import traceback
                traceback.print_exc()
                return f"<h2>Error:</h2><pre>{str(e)}</pre>"

    return render_template('pdf_translate.html', languages=PDF_LANGUAGES.keys())



@app.route('/video/<path:filename>')
def serve_video(filename):
    return send_from_directory(VIDEO_FOLDER, filename)

@app.route('/audio/<path:filename>')
def serve_audio(filename):
    if os.path.exists(os.path.join(AUDIO_OUTPUT_FOLDER, filename)):
        return send_from_directory(AUDIO_OUTPUT_FOLDER, filename)
    return send_from_directory(AUDIO_ENG_FOLDER, filename)

@app.route('/image/<path:filename>')
def serve_image(filename):
    return send_from_directory(IMAGE_FOLDER, filename)

#-----------------------------------------------------------------------#
@app.route('/text-to-text', methods=['GET', 'POST'])
def text_to_text():
    translated_text = None
    source_text = None
    source_lang = None
    target_lang = None
    
    if request.method == 'POST':
        source_text = request.form.get('sourceText', '').strip()
        source_lang = request.form.get('sourceLanguage', 'en')
        target_lang = request.form.get('targetLanguage', 'ta')
        
        if source_text:
            try:
                # Use your TextToLan class for translation
                translator = TextToLan(source_text, source_lang=source_lang, target_lang=target_lang)
                translated_text = translator.convert()
                
                # Clear memory after translation
                clear_memory()
                
            except Exception as e:
                print(f"Translation Error: {e}")
                translated_text = "Error: Translation failed. Please try again."
    
    return render_template('text_to_text.html',
                         source_text=source_text,
                         translated_text=translated_text,
                         source_lang=source_lang,
                         target_lang=target_lang)


if __name__ == "__main__":
    # Threaded=False is critical for Word Automation
    app.run(debug=True, port=5000, threaded=False)